import os
from PyPDF2 import PdfReader
from docx import Document

def read_text_file(file_path):
    """Read a plain text file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
        return {
            "success": True,
            "text": content,
            "word_count": len(content.split()),
            "char_count": len(content)
        }
    except Exception as e:
        return {
            "success": False,
            "error": str(e)
        }

def read_pdf_file(file_path):
    """Read a PDF file"""
    try:
        reader = PdfReader(file_path)
        
        # Extract text from all pages
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        
        return {
            "success": True,
            "text": text,
            "pages": len(reader.pages),
            "word_count": len(text.split()),
            "char_count": len(text)
        }
    except Exception as e:
        return {
            "success": False,
            "error": str(e)
        }

def read_word_file(file_path):
    """Read a Word document (.docx)"""
    try:
        doc = Document(file_path)
        
        # Extract text from all paragraphs
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        return {
            "success": True,
            "text": text,
            "paragraphs": len(doc.paragraphs),
            "word_count": len(text.split()),
            "char_count": len(text)
        }
    except Exception as e:
        return {
            "success": False,
            "error": str(e)
        }

def read_document(file_path):
    """
    Automatically detect file type and read it
    Supports: .txt, .pdf, .docx
    """
    if not os.path.exists(file_path):
        return {
            "success": False,
            "error": f"File not found: {file_path}"
        }
    
    # Get file extension
    _, extension = os.path.splitext(file_path)
    extension = extension.lower()
    
    print(f"\n{'='*70}")
    print(f"📄 Reading: {os.path.basename(file_path)}")
    print(f"📋 Type: {extension}")
    print('='*70)
    
    # Read based on file type
    if extension == '.txt':
        result = read_text_file(file_path)
    elif extension == '.pdf':
        result = read_pdf_file(file_path)
    elif extension == '.docx':
        result = read_word_file(file_path)
    else:
        return {
            "success": False,
            "error": f"Unsupported file type: {extension}"
        }
    
    # Display results
    if result['success']:
        print(f"\n✅ Successfully read document!")
        print(f"📊 Statistics:")
        print(f"   - Characters: {result['char_count']:,}")
        print(f"   - Words: {result['word_count']:,}")
        
        if 'pages' in result:
            print(f"   - Pages: {result['pages']}")
        if 'paragraphs' in result:
            print(f"   - Paragraphs: {result['paragraphs']}")
        
        # Show preview
        preview = result['text'][:500] + "..." if len(result['text']) > 500 else result['text']
        print(f"\n📝 Preview (first 500 characters):")
        print("-"*70)
        print(preview)
        print("-"*70)
    else:
        print(f"\n❌ Error: {result['error']}")
    
    return result

def test_all_documents():
    """Test reading all documents in test_documents folder"""
    print("\n" + "="*70)
    print("          📚 DOCUMENT EXTRACTION TESTER 📚")
    print("="*70)
    
    test_folder = "test_documents"
    
    if not os.path.exists(test_folder):
        print(f"\n❌ Folder not found: {test_folder}")
        print("   Create 'test_documents' folder and add some files!")
        return
    
    # Get all files in test_documents folder
    files = [f for f in os.listdir(test_folder) if os.path.isfile(os.path.join(test_folder, f))]
    
    if not files:
        print(f"\n❌ No files found in {test_folder}/")
        return
    
    print(f"\n✅ Found {len(files)} file(s)")
    
    results = []
    
    for file_name in files:
        file_path = os.path.join(test_folder, file_name)
        result = read_document(file_path)
        
        if result['success']:
            results.append({
                "file": file_name,
                "words": result['word_count'],
                "chars": result['char_count']
            })
        
        # Ask to continue
        if file_name != files[-1]:
            input("\n👉 Press Enter to test next document...")
    
    # Summary
    if results:
        print("\n" + "="*70)
        print("📊 SUMMARY")
        print("="*70)
        print(f"Successfully processed: {len(results)}/{len(files)} files")
        print(f"\nTotal words: {sum(r['words'] for r in results):,}")
        print(f"Total characters: {sum(r['chars'] for r in results):,}")
        print("="*70)

if __name__ == "__main__":
    try:
        test_all_documents()
    except KeyboardInterrupt:
        print("\n\n⚠️  Testing interrupted")